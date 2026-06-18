"""
DOCX structured parser — outputs DocumentNode JSON array.

Each node:
  {"type":"heading","level":1,"content":"Chapter Title","order":0}
  {"type":"paragraph","content":"Body text with **bold** and *italic*","order":1}
  {"type":"table","tableData":[["Name","HP"],["Alice","11"]],"order":2}
  {"type":"image","imageUri":"/path/to/img.png","order":3}
  {"type":"quote","content":"Do not open that door","order":4}
  {"type":"list_item","content":"Clue 1: blood stains","level":0,"order":5}
"""

import json
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table


def parse_docx(file_path, image_output_dir):
    """Parse a DOCX file and return a JSON array of DocumentNode dicts."""
    doc = Document(file_path)
    nodes = []
    order = 0
    os.makedirs(image_output_dir, exist_ok=True)

    # Build image map: rId -> bytes
    image_map = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_map[rel.rId] = rel.target_part.blob

    # Calculate body font size (mode) for heading detection
    body_font_size = _calc_body_font_size(doc)

    for element in _iter_block_items(doc):
        if isinstance(element, Paragraph):
            node = _parse_paragraph(element, image_map, image_output_dir,
                                    order, body_font_size)
            if node is not None:
                nodes.append(node)
                order += 1
        elif isinstance(element, Table):
            node = _parse_table(element, order)
            if node is not None:
                nodes.append(node)
                order += 1

    return json.dumps(nodes, ensure_ascii=False)


# ---------------------------------------------------------------------------
# Block iteration (paragraphs + tables in document order)
# ---------------------------------------------------------------------------

def _iter_block_items(parent):
    """Yield Paragraph and Table objects in document order."""
    body = parent.element.body
    for child in body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


# ---------------------------------------------------------------------------
# Paragraph parsing
# ---------------------------------------------------------------------------

def _parse_paragraph(para, image_map, image_dir, order, body_font_size):
    text = para.text.strip()
    style_name = (para.style.name or "").lower() if para.style else ""

    # 1. Extract images from runs
    for run in para.runs:
        blips = run.element.findall(
            ".//%s" % qn("a:blip")
        )
        for blip in blips:
            rId = blip.get(qn("r:embed"))
            if rId and rId in image_map:
                img_name = "img_%d_%s.png" % (order, rId)
                img_path = os.path.join(image_dir, img_name)
                if not os.path.exists(img_path):
                    with open(img_path, "wb") as f:
                        f.write(image_map[rId])
                return {"type": "image", "imageUri": img_path, "order": order}

    if not text:
        return None

    # 2. Heading detection (3-layer)
    heading_level = _detect_heading(para, style_name, body_font_size)
    if heading_level > 0:
        return {
            "type": "heading",
            "level": heading_level,
            "content": text,
            "order": order,
        }

    # 3. List item
    numPr = para._element.find(qn("w:numPr"))
    if numPr is not None:
        ilvl_el = numPr.find(qn("w:ilvl"))
        level = (
            int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0
        )
        return {
            "type": "list_item",
            "content": text,
            "level": level,
            "order": order,
        }

    # 4. Quote block
    if "quote" in style_name or "引用" in style_name:  # 引用
        return {"type": "quote", "content": text, "order": order}

    # 5. Normal paragraph (with inline formatting)
    rich_text = _extract_inline_formatting(para)
    return {"type": "paragraph", "content": rich_text, "order": order}


# ---------------------------------------------------------------------------
# Heading detection (3-layer scoring)
# ---------------------------------------------------------------------------

def _detect_heading(para, style_name, body_font_size):
    """Return heading level 1-6, or 0 if not a heading."""

    # Layer 1: Word style name
    if style_name.startswith("heading"):
        try:
            num = int(style_name.replace("heading", "").strip())
            if 1 <= num <= 6:
                return num
        except ValueError:
            pass
    if "标题" in style_name:  # 标题
        for i in range(1, 7):
            if str(i) in style_name:
                return i

    text = para.text.strip()
    if not text:
        return 0

    # Layer 2: Numbering patterns
    # Chinese chapter: 第一章, 第2节
    if re.match(r"^第[一二三四五六七八九十百千\d]+[章节部篇幕]", text):
        return 1
    # Chinese numbered: 一、二、
    if re.match(r"^[一二三四五六七八九十]+[、.．]", text):
        return 2
    # Chinese parenthesized: （一）（二）
    if re.match(r"^（[一二三四五六七八九十]+）", text):
        return 3
    # English chapter: Chapter 1, Section 2
    if re.match(r"^(Chapter|Section|Part)\s+\d+", text, re.IGNORECASE):
        return 1
    # Numbered: 1. 2.
    if re.match(r"^\d+\.\s+", text):
        return 2
    # Sub-numbered: 1.1. 2.3.
    if re.match(r"^\d+\.\d+\.\s+", text):
        return 3
    # Parenthesized: (1) (2)
    if re.match(r"^\(\d+\)", text):
        return 3
    # Chinese parenthesized: (1) (2)
    if re.match(r"^（\d+）", text):
        return 3

    # Layer 3: Typography (font size + bold + center)
    max_size = 0.0
    all_bold = True
    has_runs = False
    for run in para.runs:
        if run.text and run.text.strip():
            has_runs = True
            if run.font.size:
                pt = run.font.size.pt
                if pt > max_size:
                    max_size = pt
            if not run.bold:
                all_bold = False

    if not has_runs:
        return 0

    is_center = para.alignment in (
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    diff = max_size - body_font_size if max_size > 0 else 0

    if diff >= 6 and is_center:
        return 1
    if diff >= 4 and all_bold:
        return 2
    if diff >= 2 and all_bold and is_center:
        return 3

    return 0


# ---------------------------------------------------------------------------
# Body font size calculation (mode)
# ---------------------------------------------------------------------------

def _calc_body_font_size(doc):
    """Calculate the most common font size in the document (body text baseline)."""
    sizes = {}
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size:
                pt = round(run.font.size.pt, 1)
                sizes[pt] = sizes.get(pt, 0) + 1
    if not sizes:
        return 12.0
    return max(sizes, key=sizes.get)


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------

def _extract_inline_formatting(para):
    """Extract inline formatting as lightweight markdown markers."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            parts.append("***%s***" % text)
        elif run.bold:
            parts.append("**%s**" % text)
        elif run.italic:
            parts.append("*%s*" % text)
        elif run.font.strike:
            parts.append("~~%s~~" % text)
        else:
            parts.append(text)
    return "".join(parts)


# ---------------------------------------------------------------------------
# Table parsing
# ---------------------------------------------------------------------------

def _parse_table(table, order):
    """Extract table as a 2D array."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # Join multiple paragraphs in a cell
            cell_text = " ".join(
                p.text.strip() for p in cell.paragraphs if p.text.strip()
            )
            cells.append(cell_text)
        rows.append(cells)
    if not rows:
        return None
    return {"type": "table", "tableData": rows, "order": order}
